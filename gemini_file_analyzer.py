"""Gemini-backed multimodal source analysis with deterministic local fallback."""

from __future__ import annotations

import base64
import io
import json
import mimetypes
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Tuple

import requests
import streamlit as st

import model_router


@dataclass
class FileAnalysisResult:
    text: str
    provider: str
    model: str
    status: str
    diagnostic: str = ""


MULTIMODAL_SUFFIXES = {
    ".pdf",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".mp3",
    ".wav",
    ".m4a",
    ".mp4",
    ".mov",
}
TEXT_SUFFIXES = {".txt", ".md", ".csv", ".json"}


def _secret(name: str, default: str = "") -> str:
    try:
        return str(st.secrets.get(name, default))
    except Exception:
        return default


def _as_bool(name: str, default: bool = False) -> bool:
    raw = _secret(name, "true" if default else "false").strip().lower()
    return raw in {"1", "true", "yes", "on"}


def provider_status() -> Dict[str, Any]:
    status = model_router.provider_status("file_analysis")
    return {
        **status,
        "mode": "multimodal" if status["provider"] == "gemini" and status["available"] else "local-fallback",
    }


def _language_instruction(language: str) -> str:
    clean = (language or "English").strip().lower()
    if clean in {"arabic", "ar", "العربية"}:
        return "Write the full analysis in clear Modern Standard Arabic. Keep API names and code identifiers unchanged."
    if clean in {"french", "fr", "français", "francais"}:
        return "Write the full analysis in clear academic French. Keep API names and code identifiers unchanged."
    return "Write the full analysis in clear academic English."


def _analysis_prompt(name: str, project_context: Dict[str, Any], language: str) -> str:
    compact_context = {
        "domain": project_context.get("domain", ""),
        "program": project_context.get("program_name", ""),
        "unit": project_context.get("unit_title", ""),
        "target_concept": project_context.get("target_concept", ""),
        "target_learners": project_context.get("target_learners", ""),
        "level": project_context.get("learner_level", ""),
        "teaching_preferences": project_context.get("teaching_preferences", ""),
        "assessment_preferences": project_context.get("assessment_preferences", ""),
    }
    return f"""
You are the 3alimnIA source-analysis engine. Analyze the attached teacher-supplied file named: {name}

Project context:
{json.dumps(compact_context, ensure_ascii=False, indent=2)}

Produce a source-grounded educational analysis with these headings:
1. Source identity and apparent purpose
2. Verified concepts, definitions, and claims found in the file
3. Learning objectives or competencies supported by the file
4. Prerequisites and concept dependencies
5. Examples, activities, assessment ideas, and teaching strategies present in the file
6. Tables, diagrams, figures, screenshots, or visual evidence and what each contributes
7. Terminology and multilingual translation cautions
8. Potential misconceptions, ambiguities, missing evidence, or outdated information
9. Recommended use inside the 3alimnIA unit
10. Traceable source notes: page, section, scene, or timestamp whenever visible

Rules:
- Analyze only what is supported by the attached file.
- Never invent page numbers, quotations, references, or scientific facts.
- Distinguish source content from your pedagogical recommendation.
- If text or a visual is unreadable, say so explicitly.
- Do not include personal data that is irrelevant to teaching.
- Keep the result implementation-ready and concise enough to embed in a project prompt.
- {_language_instruction(language)}
""".strip()


def _mime_type(name: str, supplied: str = "") -> str:
    if supplied and supplied != "application/octet-stream":
        return supplied
    guessed, _ = mimetypes.guess_type(name)
    if guessed:
        return guessed
    suffix = Path(name).suffix.lower()
    return {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".md": "text/markdown",
        ".json": "application/json",
        ".m4a": "audio/mp4",
        ".mov": "video/quicktime",
    }.get(suffix, "application/octet-stream")


def _response_text(data: Dict[str, Any]) -> str:
    candidates = data.get("candidates", [])
    parts = candidates[0].get("content", {}).get("parts", []) if candidates else []
    return "\n".join(str(part.get("text") or "") for part in parts if part.get("text")).strip()


def analyze_bytes_with_gemini(
    *,
    name: str,
    raw: bytes,
    mime_type: str,
    project_context: Dict[str, Any],
    language: str,
) -> FileAnalysisResult:
    selections = [item for item in model_router.generation_candidates("file_analysis") if item.provider == "gemini"]
    if not selections or not selections[0].available:
        return FileAnalysisResult("", "local", "local-extractor", "not_configured", "GEMINI_API_KEY is not configured.")

    selection = selections[0]
    encoded = base64.b64encode(raw).decode("ascii")
    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [
                    {"text": _analysis_prompt(name, project_context, language)},
                    {"inlineData": {"mimeType": mime_type, "data": encoded}},
                ],
            }
        ],
        "generationConfig": {
            "temperature": 0.1,
            "maxOutputTokens": int(_secret("FILE_ANALYSIS_MAX_OUTPUT_TOKENS", "6000") or "6000"),
        },
    }
    url = f"{selection.base_url}/models/{selection.model}:generateContent"
    try:
        response = requests.post(
            url,
            headers={"x-goog-api-key": selection.api_key, "Content-Type": "application/json"},
            json=payload,
            timeout=int(_secret("FILE_ANALYSIS_TIMEOUT_SECONDS", "120") or "120"),
        )
        if response.status_code != 200:
            raise RuntimeError(f"Gemini HTTP {response.status_code}: {response.text[:1200]}")
        data = response.json()
        text = _response_text(data)
        if not text:
            raise RuntimeError(f"Gemini returned an empty file analysis: {json.dumps(data)[:800]}")
        return FileAnalysisResult(text, "gemini", selection.model, "completed")
    except Exception as exc:
        return FileAnalysisResult("", "gemini", selection.model, "error", str(exc))


def analyze_text_with_gemini(
    *,
    name: str,
    text: str,
    project_context: Dict[str, Any],
    language: str,
) -> FileAnalysisResult:
    selections = [item for item in model_router.generation_candidates("file_analysis") if item.provider == "gemini"]
    if not selections or not selections[0].available:
        return FileAnalysisResult("", "local", "local-extractor", "not_configured", "GEMINI_API_KEY is not configured.")
    selection = selections[0]
    prompt = _analysis_prompt(name, project_context, language) + "\n\nExtracted source text:\n" + text[:70000]
    payload = {
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0.1, "maxOutputTokens": 5000},
    }
    try:
        response = requests.post(
            f"{selection.base_url}/models/{selection.model}:generateContent",
            headers={"x-goog-api-key": selection.api_key, "Content-Type": "application/json"},
            json=payload,
            timeout=int(_secret("FILE_ANALYSIS_TIMEOUT_SECONDS", "120") or "120"),
        )
        if response.status_code != 200:
            raise RuntimeError(f"Gemini HTTP {response.status_code}: {response.text[:1200]}")
        data = response.json()
        analyzed = _response_text(data)
        if not analyzed:
            raise RuntimeError("Gemini returned an empty text analysis.")
        return FileAnalysisResult(analyzed, "gemini", selection.model, "completed")
    except Exception as exc:
        return FileAnalysisResult("", "gemini", selection.model, "error", str(exc))


def _local_extract(name: str, raw: bytes) -> Tuple[str, str]:
    suffix = Path(name).suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return raw.decode("utf-8", errors="replace"), "local-text"
    if suffix == ".docx":
        from docx import Document

        document = Document(io.BytesIO(raw))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(paragraphs), "local-docx"
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw))
        pages = []
        for index, page in enumerate(reader.pages[:80], start=1):
            page_text = (page.extract_text() or "").strip()
            if page_text:
                pages.append(f"[Page {index}]\n{page_text}")
        return "\n\n".join(pages), "local-pdf"
    return "", "unsupported"


def extract_uploaded_sources(
    uploaded_files: Any,
    *,
    project_context: Dict[str, Any],
    language: str,
) -> str:
    """Analyze uploaded teacher sources with Gemini and deterministic fallbacks."""
    if not uploaded_files:
        return ""

    max_files = max(1, int(_secret("FILE_ANALYSIS_MAX_FILES", "6") or "6"))
    max_file_mb = max(1, int(_secret("FILE_ANALYSIS_MAX_FILE_MB", "12") or "12"))
    max_total_chars = max(10000, int(_secret("FILE_ANALYSIS_MAX_TOTAL_CHARS", "90000") or "90000"))
    use_gemini = provider_status()["mode"] == "multimodal"
    enrich_text = _as_bool("GEMINI_ANALYZE_TEXT_FILES", True)

    chunks: List[str] = []
    total_chars = 0
    for uploaded in list(uploaded_files)[:max_files]:
        name = str(getattr(uploaded, "name", "source")).strip() or "source"
        suffix = Path(name).suffix.lower()
        try:
            raw = uploaded.getvalue()
        except Exception as exc:
            chunks.append(f"\n[Could not read {name}: {exc}]\n")
            continue

        if len(raw) > max_file_mb * 1024 * 1024:
            chunks.append(f"\n[Skipped {name}: file exceeds {max_file_mb} MB]\n")
            continue

        result: FileAnalysisResult | None = None
        local_text = ""
        local_method = ""

        if use_gemini and suffix in MULTIMODAL_SUFFIXES:
            result = analyze_bytes_with_gemini(
                name=name,
                raw=raw,
                mime_type=_mime_type(name, str(getattr(uploaded, "type", "") or "")),
                project_context=project_context,
                language=language,
            )

        if not result or result.status != "completed":
            try:
                local_text, local_method = _local_extract(name, raw)
            except Exception as exc:
                local_text = ""
                local_method = f"local-error: {exc}"

            if use_gemini and local_text and enrich_text:
                enriched = analyze_text_with_gemini(
                    name=name,
                    text=local_text,
                    project_context=project_context,
                    language=language,
                )
                if enriched.status == "completed":
                    result = enriched

        if result and result.status == "completed":
            text_value = result.text.strip()
            provenance = f"provider={result.provider}; model={result.model}; mode=multimodal-analysis"
        elif local_text.strip():
            text_value = local_text.strip()
            provenance = f"provider=local; model={local_method}; mode=text-extraction"
            if result and result.diagnostic:
                provenance += f"; gemini_fallback_reason={result.diagnostic[:240]}"
        else:
            diagnostic = result.diagnostic if result else "Unsupported file type or no readable content."
            text_value = f"[No usable analysis for {name}: {diagnostic}]"
            provenance = "provider=none; mode=unavailable"

        available = max_total_chars - total_chars
        if available <= 0:
            break
        text_value = text_value[:available]
        chunk = f"\n\n--- Uploaded source: {name} ---\n[{provenance}]\n{text_value}"
        chunks.append(chunk)
        total_chars += len(text_value)

    return "".join(chunks).strip()
